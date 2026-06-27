import time
import os, subprocess, pywinauto
import shutil, pyautogui, pyperclip
import sqlite3, json
import config
import datetime
from google.genai import types
from assets.local import local_path
from .client import client
from pathlib import Path
from docx import Document
from openpyxl import Workbook
from pptx import Presentation
import pypdf, docx, openpyxl, pptx
import psutil
import mss
from ddgs import DDGS
import requests
from bs4 import BeautifulSoup

permissions = {
    "open_application":"SAFE",
    "focus_application":"SAFE",
    "resize_application":"SAFE",
    "close_application":"SENSITIVE",
    "browser_new_tab":"SAFE",
    "browser_close_tab":"SENSITIVE",
    "search_browser":"SENSITIVE",
    "create_file":"SAFE",
    "create_folder":"SAFE",
    "delete_file":"SENSITIVE",
    "move_copy_file":"SENSITIVE",
    "rename_file":"SAFE",
    "search_files":"SAFE",
    "read_file":"SAFE",
    "write_file":"SENSITIVE",
    "get_system_status":"SAFE",
    "volume_control":"SAFE",
    "brightness_control":"SAFE",
    "list_processes":"SAFE",
    "clipboard":"SENSITIVE",
    "screenshot":"SENSITIVE",
    "execute_file":"SENSITIVE",
    "get_time":"SAFE",
    "web_search":"SAFE",
    "scrape_page":"SAFE"
}


conn = sqlite3.connect(config.ACTION_DB)



def init_action_log():
    cur = conn.cursor()
    cur.execute('''CREATE TABLE IF NOT EXISTS actions(
                id INTEGER PRIMARY KEY,
                timestamp TEXT NOT NULL,
                action_name TEXT NOT NULL,
                parameters TEXT NOT NULL,
                permission_level TEXT NOT NULL,
                status TEXT NOT NULL,
                result_or_error TEXT
                )''')
    
    
    conn.commit()

def log_action(action_name, parameters, permission_level, status, result_or_error):
    timestamp = datetime.datetime.now().isoformat()
    param_json = json.dumps(parameters)

    cur = conn.cursor()
    cur.execute("INSERT INTO actions(timestamp, action_name, parameters, permission_level, status, result_or_error) VALUES(?,?,?,?,?,?)", (timestamp, action_name, param_json, permission_level, status, result_or_error))
    conn.commit()


def _is_path_allowed(path, allowed_paths):
    abs_path = os.path.abspath(path)
    for dir in allowed_paths:
        if os.path.commonpath([abs_path,dir]) == dir:
            return True
    return False


# App Control
def _open_application(app_name):
    try:
        for app_path_name in local_path.APP_PATHS:
            if app_name.lower() in app_path_name:
                app_path = local_path.APP_PATHS[app_path_name]
                if "!" in app_path:
                    subprocess.Popen(["explorer.exe", f"shell:AppsFolder\\{app_path}"])
                else:
                    subprocess.Popen(local_path.APP_PATHS[app_path_name.lower()])
                return ("success", "Application opened!!")
        which_path = shutil.which(app_name.lower())
        if which_path:
            subprocess.Popen(which_path)
            return ("success", "Application opened!!")
        else:
            return ("fail", "App path not found!!")
        
    except Exception as e:
        return ("fail", "Error..." + str(e))

def _focus_application(app_name):
    try:
        app = pywinauto.Application(backend = "win32").connect(title_re = f".*{app_name}.*",found_index=0)
        time.sleep(0.5)
        app.top_window().set_focus()
        return ("success", f"{app_name} focused.")
    except Exception as e:
        return ("fail", str(e))
    
def _resize_application(app_name, action):
    try:
        search_title = local_path.WINDOW_TITLE_ALIASES.get(app_name.lower(), app_name)
        app = pywinauto.Application(backend="win32").connect(title_re=f".*{search_title}.*", found_index=0)
        time.sleep(0.5)
        if action.lower() == "minimize":
            app.top_window().set_focus()
            time.sleep(0.3)
            pyautogui.hotkey('win', 'down')
            return ("success", f"{app_name} minimized.")

        elif action.lower() == "maximize":
            app.top_window().set_focus()
            time.sleep(0.3)
            pyautogui.hotkey('win', 'up')
            return ("success", f"{app_name} maximized.")
        
        elif action.lower() == "left":
            app.top_window().set_focus()
            time.sleep(0.3)
            pyautogui.hotkey('win', 'left')
            return ("success", f"{app_name} snapped to left.")

        elif action.lower() == "right":
            app.top_window().set_focus()
            time.sleep(0.3)
            pyautogui.hotkey('win', 'right')
            return ("success", f"{app_name} snapped to right.")
        else:
            return ("fail", "invalid action")
    except Exception as e:
        return ("fail", str(e))

def _close_application(app_name):
    try:
        app = pywinauto.Application(backend = "win32").connect(title_re = f".*{app_name}.*",found_index=0)
        # time.sleep(0.5)
        app.top_window().close()
        return ("success", f"{app_name} closed.")
    except Exception as e:
        return ("fail", str(e))

def _browser_new_tab():
    try:
        app = pywinauto.Application(backend="win32").connect(title_re=f".*{local_path.DEFAULT_BROWSER}.*", found_index=0)
        app.top_window().set_focus()
        pyautogui.hotkey('ctrl', 't')
        return ("success", "New tab opened.")
    except Exception as e:
        try:
            _open_application(local_path.DEFAULT_BROWSER)
            time.sleep(2)
            app = pywinauto.Application(backend="win32").connect(title_re=f".*{local_path.DEFAULT_BROWSER}.*", found_index=0)
            app.top_window().set_focus()
            pyautogui.hotkey('ctrl', 't')
            return ("success", "New tab opened.")
        except Exception as e2:
            return ("fail", str(e2))
    
def _browser_close_tab():
    try:
        app = pywinauto.Application(backend="win32").connect(title_re=f".*{local_path.DEFAULT_BROWSER}.*", found_index=0)
        app.top_window().set_focus()
        pyautogui.hotkey('ctrl', 'w')
        return ("success", "Tab closed.")
    except Exception as e:
        return ("fail", str(e))
    
def _search_browser(query):
    try:
        _browser_new_tab()
        app = pywinauto.Application(backend="win32").connect(title_re=f".*{local_path.DEFAULT_BROWSER}.*", found_index=0)
        app.top_window().set_focus()
        time.sleep(0.5)
        pyautogui.hotkey('ctrl', 'l')
        time.sleep(0.5)
        pyperclip.copy(query) 
        pyautogui.hotkey('ctrl', 'v')
        pyautogui.press('enter')
        return ("success", "Navigated successfully.")
    except Exception as e:
        return ("fail", str(e))
    
def _create_file(filename, file_type):
    try:
        file_type = file_type.lower()
        
        if file_type == "word":
            full_path = os.path.join(local_path.ALLOWED_WRITE_PATHS[0], filename + ".docx")
            if not _is_path_allowed(full_path, local_path.ALLOWED_WRITE_PATHS):
                return ("fail", "Path not allowed.")
            doc = Document()
            doc.save(full_path)
            
        elif file_type == "excel":
            full_path = os.path.join(local_path.ALLOWED_WRITE_PATHS[0], filename + ".xlsx")
            if not _is_path_allowed(full_path, local_path.ALLOWED_WRITE_PATHS):
                return ("fail", "Path not allowed.")
            wb = Workbook()
            wb.save(full_path)
            
        elif file_type == "powerpoint":
            full_path = os.path.join(local_path.ALLOWED_WRITE_PATHS[0], filename + ".pptx")
            if not _is_path_allowed(full_path, local_path.ALLOWED_WRITE_PATHS):
                return ("fail", "Path not allowed.")
            prs = Presentation()
            prs.save(full_path)
            
        else:
            full_path = os.path.join(local_path.ALLOWED_WRITE_PATHS[0], filename + "." + file_type)
            if not _is_path_allowed(full_path, local_path.ALLOWED_WRITE_PATHS):
                return ("fail", "Path not allowed.")
            Path(full_path).write_text("")
            
        return ("success", f"{filename} created at {full_path}")
    except Exception as e:
        return ("fail", str(e))


def _create_folder(folder_name):
    try:
        path = os.path.join(local_path.ALLOWED_WRITE_PATHS[0], folder_name)
        if not _is_path_allowed(path, local_path.ALLOWED_WRITE_PATHS):
            return ("fail", "Path not allowed.")
        os.makedirs(path, exist_ok=True)
        return ("success", f"{folder_name} folder created.")
    except Exception as e:
        return ("fail", str(e))
    

def _delete_file(file_name, type, file_type=""):
    try:
        if type == "file":
            path = os.path.join(local_path.ALLOWED_WRITE_PATHS[0], f'{file_name}.{file_type}')
        elif type == "folder":
            path = os.path.join(local_path.ALLOWED_WRITE_PATHS[0], file_name)
        else:
            return ("fail", "Invalid type. Must be 'file' or 'folder'.")
        
        if not _is_path_allowed(path, local_path.ALLOWED_WRITE_PATHS):
            return ("fail", "Path not allowed.")
        
        if type == "file":
            if os.path.isfile(path):
                os.remove(path)
                return ("success", f"{file_name}.{file_type} deleted.")
            else:
                return ("fail", "File not found.")
        elif type == "folder":
            if os.path.isdir(path):
                shutil.rmtree(path)
                return ("success", f"{file_name} folder deleted.")
            else:
                return ("fail", "Folder not found.")
    except Exception as e:
        return ("fail", str(e))
    

def _move_copy_file(file_name, destination, operation, type, file_type=""):
    try:
        if type == "file":
            source_path = os.path.join(local_path.ALLOWED_WRITE_PATHS[0], f'{file_name}.{file_type}')
        elif type == "folder":
            source_path = os.path.join(local_path.ALLOWED_WRITE_PATHS[0], file_name)
        
        dest_path = os.path.join(local_path.ALLOWED_WRITE_PATHS[0], destination)
        
        if not _is_path_allowed(source_path, local_path.ALLOWED_WRITE_PATHS):
            return ("fail", "Source path not allowed.")
        if not _is_path_allowed(dest_path, local_path.ALLOWED_WRITE_PATHS):
            return ("fail", "Destination path not allowed.")
        
        if operation == "move":
            shutil.move(source_path, dest_path)
            return ("success", f"{file_name} moved to {destination}.")
        elif operation == "copy":
            if type == "file":
                shutil.copy(source_path, dest_path)
            elif type == "folder":
                shutil.copytree(source_path, dest_path)
            return ("success", f"{file_name} copied to {destination}.")
        else:
            return ("fail", "Invalid operation. Must be 'move' or 'copy'.")
    except Exception as e:
        return ("fail", str(e))
    
def _rename_file(file_name, new_name, type, file_type=""):
    try:
        if type == "file":
            source_path = os.path.join(local_path.ALLOWED_WRITE_PATHS[0], f'{file_name}.{file_type}')
            dest_path = os.path.join(local_path.ALLOWED_WRITE_PATHS[0], f'{new_name}.{file_type}')
        elif type == "folder":
            source_path = os.path.join(local_path.ALLOWED_WRITE_PATHS[0], file_name)
            dest_path = os.path.join(local_path.ALLOWED_WRITE_PATHS[0], new_name)
        else:
            return ("fail", "Invalid type. Must be 'file' or 'folder'.")
        
        if not _is_path_allowed(source_path, local_path.ALLOWED_WRITE_PATHS):
            return ("fail", "Source path not allowed.")
        if not _is_path_allowed(dest_path, local_path.ALLOWED_WRITE_PATHS):
            return ("fail", "Destination path not allowed.")
        
        os.rename(source_path, dest_path)
        return ("success", f"{file_name} renamed to {new_name}.")
    except Exception as e:
        return ("fail", str(e))

def _search_files(query, search_by):
    try:
        results = []
        for dirpath, dirnames, files in os.walk(local_path.ALLOWED_WRITE_PATHS[0]):
            if search_by == "name":
                for file in files:
                    if query.lower() in file.lower():
                        results.append(os.path.join(dirpath, file))
                for dirname in dirnames:
                    if query.lower() in dirname.lower():
                        results.append(os.path.join(dirpath, dirname))
            
            elif search_by == "type":
                for file in files:
                    if file.lower().endswith(f".{query.lower()}"):
                        results.append(os.path.join(dirpath, file))
            
            elif search_by == "content":
                for file in files:
                    file_path = os.path.join(dirpath, file)
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            if query.lower() in f.read().lower():
                                results.append(file_path)
                    except:
                        pass
            else:
                return ("fail", "Invalid search_by. Must be 'name', 'type', or 'content'.")
        
        if results:
            return ("success", f"Found {len(results)} result(s):\n" + "\n".join(results))
        else:
            return ("success", "No results found.")
    except Exception as e:
        return ("fail", str(e))
    

def _read_file(file_name, file_type):
    try:
        path = os.path.join(local_path.ALLOWED_WRITE_PATHS[0], f'{file_name}.{file_type}')
        
        if not _is_path_allowed(path, local_path.ALLOWED_WRITE_PATHS):
            return ("fail", "Path not allowed.")
        
        if file_type == "docx":
            doc = Document(path)
            text = "\n".join([para.text for para in doc.paragraphs])
            
        elif file_type == "xlsx":
            wb = openpyxl.load_workbook(path)
            text = ""
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                for row in ws.iter_rows(values_only=True):
                    text += " | ".join([str(cell) for cell in row if cell is not None]) + "\n"
                    
        elif file_type == "pdf":
            reader = pypdf.PdfReader(path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
                
        else:
            with open(path, 'r', encoding='utf-8') as f:
                text = f.read()
        
        return ("success", text)
    except Exception as e:
        return ("fail", str(e))

def _write_file(file_name, file_type, content, mode="write"):
    print(f"_write_file called: file={file_name}.{file_type}, mode={mode}, content={content}")
    try:
        path = os.path.join(local_path.ALLOWED_WRITE_PATHS[0], f'{file_name}.{file_type}')
        
        if not _is_path_allowed(path, local_path.ALLOWED_WRITE_PATHS):
            return ("fail", "Path not allowed.")
        
        if file_type == "docx":
            if mode == "write" or not os.path.exists(path):
                doc = Document()
            else:
                doc = Document(path)
                print(f"Existing paragraphs: {len(doc.paragraphs)}")
                doc.add_paragraph(content)
                print(f"Paragraphs after add: {len(doc.paragraphs)}")
                doc.save(path)
            
            
        elif file_type == "xlsx":
            if os.path.exists(path):
                wb = openpyxl.load_workbook(path)
                ws = wb.active
            else:
                wb = Workbook()
                ws = wb.active
            if mode == "write":
                ws.delete_rows(1, ws.max_row)
            next_row = ws.max_row + 1
            ws.append([content])
            wb.save(path)
            
        elif file_type == "pdf":
            return ("fail", "PDF files cannot be edited directly.")
        
        else:
            file_mode = 'w' if mode == "write" else 'a'
            with open(path, file_mode, encoding='utf-8') as f:
                f.write(content)
        
        return ("success", f"{file_name}.{file_type} updated.")
    except Exception as e:
        return ("fail", str(e))
    
def _get_system_status(metric="all"):
    try:
        if metric == "cpu":
            cpu = psutil.cpu_percent(interval=1)
            return ("success", f"CPU usage is at {cpu} percent.")
        
        elif metric == "ram":
            ram = psutil.virtual_memory()
            used = round(ram.used / (1024**3), 1)
            total = round(ram.total / (1024**3), 1)
            return ("success", f"RAM usage is {used}GB out of {total}GB, at {ram.percent} percent.")
        
        elif metric == "battery":
            battery = psutil.battery()
            if battery is None:
                return ("success", "No battery detected — running on desktop power.")
            status = "charging" if battery.power_plugged else "not charging"
            return ("success", f"Battery at {round(battery.percent)} percent, {status}.")
        
        elif metric == "network":
            net = psutil.net_io_counters()
            sent = round(net.bytes_sent / (1024**2), 1)
            recv = round(net.bytes_recv / (1024**2), 1)
            return ("success", f"Network — {sent}MB sent, {recv}MB received since last boot.")
        
        elif metric == "temperature":
            temps = psutil.sensors_temperatures()
            if not temps:
                return ("success", "Temperature sensors not available on this system.")
            readings = []
            for name, entries in temps.items():
                for entry in entries:
                    readings.append(f"{entry.label or name}: {entry.current}°C")
            return ("success", "Temperatures — " + ", ".join(readings))
        
        elif metric == "all":
            cpu = psutil.cpu_percent(interval=1)
            ram = psutil.virtual_memory()
            battery = psutil.battery()
            net = psutil.net_io_counters()
            
            ram_used = round(ram.used / (1024**3), 1)
            ram_total = round(ram.total / (1024**3), 1)
            net_sent = round(net.bytes_sent / (1024**2), 1)
            net_recv = round(net.bytes_recv / (1024**2), 1)
            
            summary = f"CPU at {cpu}%. RAM {ram_used}GB of {ram_total}GB ({ram.percent}%). "
            
            if battery:
                status = "charging" if battery.power_plugged else "not charging"
                summary += f"Battery at {round(battery.percent)}%, {status}. "
            
            summary += f"Network — {net_sent}MB sent, {net_recv}MB received."
            return ("success", summary)
        
        else:
            return ("fail", "Invalid metric. Use cpu, ram, battery, network, temperature, or all.")
    
    except Exception as e:
        return ("fail", str(e))
    
def _volume_control(action, level=None):
    try:
        from pycaw.pycaw import AudioUtilities

        device = AudioUtilities.GetSpeakers()
        volume = device.EndpointVolume

        if action == "get":
            current = round(volume.GetMasterVolumeLevelScalar() * 100)
            return ("success", f"Volume is at {current} percent.")
        elif action == "set":
            volume.SetMasterVolumeLevelScalar(level / 100, None)
            return ("success", f"Volume set to {level} percent.")
        else:
            return ("fail", "Invalid action. Must be 'get' or 'set'.")
    except Exception as e:
        print(f"Volume error: {str(e)}")
        return ("fail", str(e))

def _brightness_control(action, level=None):
    try:
        import screen_brightness_control as sbc
        
        if action == "get":
            current = sbc.get_brightness()[0]
            return ("success", f"Screen brightness is at {current} percent.")
        elif action == "set":
            sbc.set_brightness(level)
            return ("success", f"Brightness set to {level} percent.")
        else:
            return ("fail", "Invalid action. Must be 'get' or 'set'.")
    except Exception as e:
        return ("fail", str(e))
    
def _list_processes():
    try:
        processes = []
        for proc in psutil.process_iter(['pid', 'name']):
            try:
                processes.append(f"{proc.info['name']} (PID: {proc.info['pid']})")
            except (psutil.NoSuchProcess, psutil.AccessDenied):
                pass
        
        if processes:
            return ("success", f"Running processes:\n" + "\n".join(processes))
        else:
            return ("success", "No processes found.")
    except Exception as e:
        return ("fail", str(e))
    
def _clipboard(action, text=""):
    try:
        if action == "get":
            content = pyperclip.paste()
            return ("success", f"Clipboard contains: {content}")
        elif action == "set":
            pyperclip.copy(text)
            return ("success", "Clipboard updated.")
        else:
            return ("fail", "Invalid action. Must be 'get' or 'set'.")
    except Exception as e:
        return ("fail", str(e))


def _screenshot(mode="full", app_name=""):
    try:
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"screenshot_{timestamp}.png"
        filepath = os.path.join(local_path.ALLOWED_WRITE_PATHS[0], filename)

        if mode == "full":
            with mss.mss() as sct:
                sct.shot(output=filepath)

        elif mode == "window":
            app = pywinauto.Application(backend="win32").connect(
                title_re=f".*{app_name}.*", found_index=0)
            rect = app.top_window().rectangle()
            with mss.mss() as sct:
                monitor = {
                    "top": rect.top,
                    "left": rect.left,
                    "width": rect.width(),
                    "height": rect.height()
                }
                sct.shot(mon=monitor, output=filepath)
        else:
            return ("fail", "Invalid mode. Must be 'full' or 'window'.")

        return ("success", f"Screenshot saved to {filepath}")
    except Exception as e:
        return ("fail", str(e)) 
    
def _execute_file(filename):
    try:
        filepath = os.path.join(local_path.ALLOWED_WRITE_PATHS[0], filename)
        
        if not _is_path_allowed(filepath, local_path.ALLOWED_WRITE_PATHS):
            return ("fail", "Path not allowed.")
        
        if not os.path.exists(filepath):
            return ("fail", f"{filename} not found in workspace folder.")
        
        ext = os.path.splitext(filename)[1].lower()
        
        if ext == ".py":
            subprocess.Popen(["python", filepath])
        elif ext == ".js":
            subprocess.Popen(["node", filepath])
        elif ext in [".bat", ".exe"]:
            subprocess.Popen([filepath])
        else:
            return ("fail", f"File type {ext} is not supported for execution.")
        
        return ("success", f"{filename} is now running.")
    except Exception as e:
        return ("fail", str(e))
    

def _get_time():
    t = datetime.datetime.now().strftime("%d/%m/%Y, %H:%M:%S")
    return ("success", f"{t} is the current time.")

def _web_search(query):
    try:
        with DDGS() as ddgs:
            results = ddgs.text(query, max_results=3)
        links = []
        final_result = ""
        for result in results:
            links.append(result['href'])
        for link in links:
            page_result = _scrape_page(link,query)
            if page_result:
                final_result += page_result
        return ("success", f"Data extracted from {links}")
    except Exception as e:
        return ("fail", str(e))

def _scrape_page(url,question):
    try:
        page_response = requests.get(url, timeout=10)
        html_content = BeautifulSoup(page_response.content,'html.parser')
        texts = html_content.get_text().split("\n")
        texts = [text.strip() for text in texts if text.strip()]
        texts = "\n".join(texts)[:8000]
        prompt = f"Answer the give question: {question}\n Using the data given:{texts}"
        response = client.models.generate_content(model=config.GEMINI_MODEL, contents= prompt)
        return response.text
    except Exception as e:
        return None
    
_actions = {
    "open_application":_open_application,
    "focus_application":_focus_application,
    "resize_application":_resize_application,
    "close_application":_close_application,
    "browser_new_tab":_browser_new_tab,
    "browser_close_tab":_browser_close_tab,
    "search_browser":_search_browser,
    "create_file":_create_file,
    "create_folder":_create_folder,
    "delete_file":_delete_file,
    "move_copy_file":_move_copy_file,
    "rename_file":_rename_file,
    "search_files":_search_files,
    "read_file":_read_file,
    "write_file":_write_file,
    "get_system_status":_get_system_status,
    "volume_control":_volume_control,
    "brightness_control":_brightness_control,
    "list_processes":_list_processes,
    "clipboard":_clipboard,
    "screenshot":_screenshot,
    "execute_file":_execute_file,
    "get_time":_get_time,
    "web_search":_web_search,
    "scrape_page":_scrape_page
}

_action_declarations = [types.FunctionDeclaration(name="open_application", 
                                                  description="Opens an application on the user's PC. Use this when the user asks to open, launch, or start any application or program.", 
                                                  parameters=types.Schema(
                                                      type=types.Type.OBJECT,
                                                      properties={
                                                          "app_name":types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The name of the application to open, exactly as it appears in the Start Menu or Desktop, without the file extension."
                                                              )
                                                      },
                                                      required=["app_name"]
                                                    )   
                                                ),
                        types.FunctionDeclaration(name="focus_application",
                                                  description="Brings an already open application window to the foreground. Use this when the user asks to focus, switch to, bring up, or show an application that is already running.",
                                                  parameters=types.Schema(
                                                      type=types.Type.OBJECT,
                                                      properties={
                                                          "app_name": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The name of the application to focus, as it would appear in the window title bar."
                                                          )
                                                      },
                                                      required=["app_name"]
                                                    )
                                                ),
                        types.FunctionDeclaration(name="resize_application",
                                                  description="Minimizes, maximizes, or snaps an open application window to the left or right half of the screen. Use this when the user asks to minimize, maximize, snap, or move an application to a side of the screen.",
                                                  parameters=types.Schema(
                                                      type=types.Type.OBJECT,
                                                      properties={
                                                          "app_name": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The name of the application as it appears in the window title bar."
                                                          ),
                                                          "action": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The action to perform. Must be exactly one of: 'minimize', 'maximize', 'left', 'right'."
                                                          )
                                                      },
                                                      required=["app_name", "action"]
                                                    )
                                                ),
                        types.FunctionDeclaration(name="close_application",
                                                  description="Closes an open application window. Use this when the user asks to close, quit, or exit an application.",
                                                  parameters=types.Schema(
                                                      type=types.Type.OBJECT,
                                                      properties={
                                                          "app_name": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The name of the application to close, as it appears in the window title bar."
                                                          )
                                                      },
                                                      required=["app_name"]
                                                  )
                                                ),
                        types.FunctionDeclaration(name="browser_new_tab",
                                                  description="Opens a new tab in the browser. Use this when the user asks to open a new tab.",
                                                  parameters=types.Schema(
                                                      type=types.Type.OBJECT,
                                                      properties={}
                                                  )
                                                ),
                        types.FunctionDeclaration(
                                                  name="browser_close_tab",
                                                  description="Closes the current tab in the browser. Use this when the user asks to close the current tab.",
                                                  parameters=types.Schema(
                                                      type=types.Type.OBJECT,
                                                      properties={}
                                                  )
                                                ),
                        types.FunctionDeclaration(name="search_browser",
                                                  description="Opens a new browser tab and navigates to a URL or performs a search. Use this when the user asks to search for something, go to a website, or open a URL.",
                                                  parameters=types.Schema(
                                                      type=types.Type.OBJECT,
                                                      properties={
                                                          "query": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The URL to navigate to or the search query to look up. Can be a full URL like 'https://github.com' or plain text like 'Python tutorials'."
                                                          )
                                                      },
                                                      required=["query"]
                                                  )
                                                ),
                        types.FunctionDeclaration(name="create_file",
                                                  description="Creates a new empty file in the ARES workspace folder. Use this when the user asks to create a file, document, spreadsheet, presentation, or code file.",
                                                  parameters=types.Schema(
                                                      type=types.Type.OBJECT,
                                                      properties={
                                                          "filename": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The name of the file to create, without extension."
                                                          ),
                                                          "file_type": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The type of file to create. Use 'word' for .docx, 'excel' for .xlsx, 'powerpoint' for .pptx, or any file extension like 'py', 'cpp', 'txt' for plain text files."
                                                          )
                                                      },
                                                      required=["filename", "file_type"]
                                                  )
                                                ),
                        types.FunctionDeclaration(name="create_folder",
                                                  description="Creates a new folder in the ARES workspace folder. Use this when the user asks to create a folder or directory.",
                                                  parameters=types.Schema(
                                                      type=types.Type.OBJECT,
                                                      properties={
                                                          "folder_name": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The name of the folder to create."
                                                          )
                                                      },
                                                      required=["folder_name"]
                                                  )
                                                ),
                        types.FunctionDeclaration(name="delete_file",
                                                  description="Deletes a file or folder from the ARES workspace folder. Use this when the user asks to delete or remove a file or folder.",
                                                  parameters=types.Schema(
                                                      type=types.Type.OBJECT,
                                                      properties={
                                                          "file_name": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The name of the file or folder to delete, without extension."
                                                          ),
                                                          "type": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="Whether to delete a file or folder. Must be exactly 'file' or 'folder'."
                                                          ),
                                                          "file_type": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The file extension without the dot, for example 'py', 'docx'. Only required when type is 'file', omit for folders."
                                                          )
                                                      },
                                                      required=["file_name", "type"]
                                                  )
                                                ),
                        types.FunctionDeclaration(name="move_copy_file",
                                                  description="Moves or copies a file or folder within the ARES workspace folder. Use this when the user asks to move, copy, or duplicate a file or folder.",
                                                  parameters=types.Schema(
                                                      type=types.Type.OBJECT,
                                                      properties={
                                                          "file_name": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The name of the file or folder to move or copy, without extension."
                                                          ),
                                                          "file_type": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The file extension without the dot, for example 'py', 'docx'. Only required when type is 'file', omit for folders."
                                                          ),
                                                          "destination": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The destination folder name within the workspace folder."
                                                          ),
                                                          "operation": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The operation to perform. Must be exactly 'move' or 'copy'."
                                                          ),
                                                          "type": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="Whether the target is a file or folder. Must be exactly 'file' or 'folder'."
                                                          )
                                                      },
                                                      required=["file_name", "destination", "operation", "type"]
                                                  )
                                                ),
                        types.FunctionDeclaration(name="rename_file",
                                                  description="Renames a file or folder in the ARES workspace folder. Use this when the user asks to rename a file or folder.",
                                                  parameters=types.Schema(
                                                      type=types.Type.OBJECT,
                                                      properties={
                                                          "file_name": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The current name of the file or folder, without extension."
                                                          ),
                                                          "new_name": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The new name for the file or folder, without extension."
                                                          ),
                                                          "type": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="Whether the target is a file or folder. Must be exactly 'file' or 'folder'."
                                                          ),
                                                          "file_type": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The file extension without the dot, for example 'py', 'docx'. Only required when type is 'file', omit for folders."
                                                          )
                                                      },
                                                      required=["file_name", "new_name", "type"]
                                                  )
                                                ),
                        types.FunctionDeclaration(name="search_files",
                                                  description="Searches for files or folders in the ARES workspace folder by name, type, or content. Use this when the user asks to find or search for a file.",
                                                  parameters=types.Schema(
                                                      type=types.Type.OBJECT,
                                                      properties={
                                                          "query": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The search term. For name search use the file name, for type search use the extension without dot like 'py', for content search use the text to find inside files."
                                                          ),
                                                          "search_by": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The search method. Must be exactly 'name', 'type', or 'content'."
                                                          )
                                                      },
                                                      required=["query", "search_by"]
                                                  )
                                                ),
                        types.FunctionDeclaration(name="read_file",
                                                  description="Reads and returns the contents of a file in the ARES workspace folder. Use this when the user asks to read, open, or get the contents of a file.",
                                                  parameters=types.Schema(
                                                      type=types.Type.OBJECT,
                                                      properties={
                                                          "file_name": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The name of the file to read, without extension."
                                                          ),
                                                          "file_type": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The file extension without the dot, for example 'py', 'docx', 'pdf', 'xlsx'."
                                                          )
                                                      },
                                                      required=["file_name", "file_type"]
                                                  )
                                                ),
                        types.FunctionDeclaration(name="write_file",
                                                  description="Writes or appends content to a file in the ARES workspace folder. Use this when the user asks to write to, edit, update, or add content to a file.",
                                                  parameters=types.Schema(
                                                      type=types.Type.OBJECT,
                                                      properties={
                                                          "file_name": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The name of the file to write to, without extension."
                                                          ),
                                                          "file_type": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The file extension without the dot, for example 'py', 'docx', 'txt'."
                                                          ),
                                                          "content": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The content to write to the file."
                                                          ),
                                                          "mode": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="Whether to overwrite the file or append to it. Must be exactly 'write' or 'append'."
                                                          )
                                                      },
                                                      required=["file_name", "file_type", "content", "mode"]
                                                  )
                                                ),
                        types.FunctionDeclaration(name="get_system_status",
                                                  description="Retrieves system status information including CPU, RAM, battery, network, and temperature. Use this when the user asks about system performance, battery level, memory usage, or hardware status.",
                                                  parameters=types.Schema(
                                                      type=types.Type.OBJECT,
                                                      properties={
                                                          "metric": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The metric to retrieve. Must be exactly one of: 'cpu', 'ram', 'battery', 'network', 'temperature', or 'all'. Defaults to 'all' if not specified."
                                                          )
                                                      },
                                                      required=[]
                                                  )
                                                ),
                        types.FunctionDeclaration(name="volume_control",
                                                  description="Gets or sets the system volume. Use this when the user asks about volume level or asks to change, increase, decrease, mute, or set the volume.",
                                                  parameters=types.Schema(
                                                      type=types.Type.OBJECT,
                                                      properties={
                                                          "action": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The action to perform. Must be exactly 'get' or 'set'."
                                                          ),
                                                          "level": types.Schema(
                                                              type=types.Type.INTEGER,
                                                              description="The volume level to set, from 0 to 100. Only required when action is 'set'."
                                                          )
                                                      },
                                                      required=["action"]
                                                  )
                                                ),
                        types.FunctionDeclaration(name="brightness_control",
                                                  description="Gets or sets the screen brightness. Use this when the user asks about brightness level or asks to change, increase, decrease, or set the screen brightness.",
                                                  parameters=types.Schema(
                                                      type=types.Type.OBJECT,
                                                      properties={
                                                          "action": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The action to perform. Must be exactly 'get' or 'set'."
                                                          ),
                                                          "level": types.Schema(
                                                              type=types.Type.INTEGER,
                                                              description="The brightness level to set, from 0 to 100. Only required when action is 'set'."
                                                          )
                                                      },
                                                      required=["action"]
                                                  )
                                                ),
                        types.FunctionDeclaration(name="clipboard",
                                                  description="Reads from or writes to the system clipboard. Use this when the user asks to copy something to clipboard, paste from clipboard, or check what is in the clipboard.",
                                                  parameters=types.Schema(
                                                      type=types.Type.OBJECT,
                                                      properties={
                                                          "action": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The action to perform. Must be exactly 'get' to read clipboard or 'set' to write to clipboard."
                                                          ),
                                                          "text": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The text to copy to clipboard. Only required when action is 'set'."
                                                          )
                                                      },
                                                      required=["action"]
                                                  )
                                                ),
                        types.FunctionDeclaration(name="screenshot",
                                                  description="Takes a screenshot of the full screen or a specific application window. Use this when the user asks to take a screenshot or capture the screen.",
                                                  parameters=types.Schema(
                                                      type=types.Type.OBJECT,
                                                      properties={
                                                          "mode": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The screenshot mode. Use 'full' for the entire screen or 'window' for a specific application window."
                                                          ),
                                                          "app_name": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The name of the application window to capture. Only required when mode is 'window'."
                                                          )
                                                      },
                                                      required=["mode"]
                                                  )
                                                ),
                        types.FunctionDeclaration(name="execute_file",
                                                  description="Executes a file from the ARES workspace folder. Use this when the user asks to run, execute, or start a script or program file.",
                                                  parameters=types.Schema(
                                                      type=types.Type.OBJECT,
                                                      properties={
                                                          "filename": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The name of the file to execute, including extension. For example 'script.py' or 'program.exe'. The file must exist in the ARES workspace folder."
                                                          )
                                                      },
                                                      required=["filename"]
                                                  )
                                                ),
                        types.FunctionDeclaration(name="get_time",
                                                  description="Returns the current time and date. Use this when the user asks what time it is, what today's date is, or what day of the week it is.",
                                                  parameters=types.Schema(
                                                      type=types.Type.OBJECT,
                                                      properties={}
                                                  )
                                                ),
                        types.FunctionDeclaration(name="web_search",
                                                  description="Searches the web and returns results for ARES to read internally. Use this only when you need current or real-time information to answer the user's question — such as live prices, recent news, or anything beyond your knowledge cutoff. Do not use this when the user asks to search or look something up — use search_browser for that instead.",
                                                  parameters=types.Schema(
                                                      type=types.Type.OBJECT,
                                                      properties={
                                                          "query": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The search query to look up."
                                                          )
                                                      },
                                                      required=["query"]
                                                  )
                                                ),
                        types.FunctionDeclaration(name="scrape_page",
                                                  description="Fetches and reads the content of a specific webpage URL and extracts information from it to answer a question. Use this after web_search returns URLs, when you need to read the actual content of a page to get specific information like prices, details, or data.",
                                                  parameters=types.Schema(
                                                      type=types.Type.OBJECT,
                                                      properties={
                                                          "url": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The full URL of the webpage to fetch and read."
                                                          ),
                                                          "question": types.Schema(
                                                              type=types.Type.STRING,
                                                              description="The specific question to answer from the page content."
                                                          )
                                                      },
                                                      required=["url", "question"]
                                                  )
                                                )
                        ]

ares_tools = types.Tool(function_declarations=_action_declarations)

def execute_action(action_name, parameters):
    if action_name in permissions:
        if permissions[action_name] == 'BLOCKED':
            log_action(action_name, parameters, permissions[action_name], 'BLOCKED', None)
            return "Action blocked..."
        elif permissions[action_name] == 'SENSITIVE':
            confirmation = input(f"Requesting permission to {action_name} with parameters {parameters}?(yes/no)")
            if confirmation.lower() == 'yes':
                func = _actions.get(action_name)
                result = func(**parameters)
                log_action(action_name, parameters, permissions[action_name], result[0], result[1])
                return result[1]
            else:
                log_action(action_name, parameters, permissions[action_name], 'DENIED', None)
                return "Permission denied..."
        elif permissions[action_name] == 'SAFE':
            func = _actions.get(action_name)
            result = func(**parameters)
            log_action(action_name, parameters, permissions[action_name], result[0], result[1])
            return result[1]
    else:
        return f"Cannot do {action_name}"
    
